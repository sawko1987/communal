import os
from edit_abonent_window import EditAbonentWindow
import customtkinter as ctk
from CTkMessagebox import CTkMessagebox
from datetime import datetime

from HistoryWindow import ConsumptionHistoryWindow
from add_abonent_window import AddAbonentWindow
from monthly_data_window import MonthlyDataWindow
from users_db import SqliteDB

# Настройка темы приложения
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

class Window:
    def __init__(self, width, height, title="Учет коммунальных услуг АО_Корммаш",
                 resizable=(False, False), icon='image/korm.ico'):
        self.root = ctk.CTk()
        self.root.title(title)
        self.root.geometry(f"{width}x{height}")
        self.root.resizable(resizable[0], resizable[1])

        # Проверяем существование файла иконки
        if icon and os.path.exists(icon):
            self.root.iconbitmap(icon)
        elif icon:
            print(f"Предупреждение: файл иконки '{icon}' не найден")

        # Инициализация базы данных
        self.db = SqliteDB()
        # Загружаем данные абонентов
        self.list_abonent = self.load_abonents()

        # Основные элементы интерфейса
        self.tab_control = None
        self.combobox = None
        self.selected_abonent_info = None

        # Инициализация интерфейса
        self.draw_widget()

        # Установим начальное значение и вызовем обработчик
        if self.list_abonent:
            self.combobox.set(self.list_abonent[0][1])
            self.on_combobox_select()

    def load_abonents(self):
        """Загружает список абонентов из базы данных"""
        try:
            abonents = self.db.fetch_data()
            print(f"Загружено абонентов: {len(abonents)}")
            return abonents
        except Exception as e:
            print(f"Ошибка при загрузке абонентов: {e}")
            return []

    def create_child_window(self, width, height, title=None):
        """Создает окно добавления абонента"""
        print("Создание дочернего окна")  # Отладочный вывод
        try:
            # Create child window without hiding main window
            child_window = AddAbonentWindow(self.root, width, height, title="Добавить абонента")
            print("Дочернее окно создано")  # Отладочный вывод
            
            # Wait for child window to close
            self.root.wait_window(child_window.root)
            print("Ожидание закрытия дочернего окна завершено")  # Отладочный вывод
            
            # Refresh data after child window closes
            self.refresh_data()
            self.on_combobox_select()
            print("Данные обновлены")  # Отладочный вывод
            
        except Exception as e:
            print(f"Ошибка при создании окна: {e}")

    def create_monthly_data_window(self, width, height, abonent_id, title=None):
        """Создает окно внесения месячных данных"""
        try:
            # Create and wait for child window without hiding main window
            monthly_window = MonthlyDataWindow(self.root, width, height, abonent_id, title=title)
            self.root.wait_window(monthly_window.root)
        except Exception as e:
            print(f"Ошибка при создании окна месячных данных: {e}")

    def create_consumption_history_window(self, width, height, abonent_id, title=None):
        """Создает окно истории потребления"""
        try:
            # Create and wait for child window without hiding main window
            history_window = ConsumptionHistoryWindow(self.root, width, height, abonent_id, title=title)
            self.root.wait_window(history_window.root)
        except Exception as e:
            print(f"Ошибка при создании окна истории: {e}")

    def draw_widget(self):
        """Создает элементы интерфейса"""
        # Создание основного контейнера
        main_container = ctk.CTkFrame(self.root)
        main_container.pack(fill="both", expand=True, padx=20, pady=20)

        # Создание контейнера с вкладками
        self.tab_control = ctk.CTkTabview(main_container, width=200)
        self.tab_control.pack(side='left', fill='y', padx=(0, 20))

        # Добавление вкладок
        self.tab_control.add("Абоненты")
        self.tab_control.add("Отчеты")
        self.tab_control.add("Настройки")

        # Получаем фреймы для каждой вкладки
        tab1 = self.tab_control.tab("Абоненты")
        tab2 = self.tab_control.tab("Отчеты")
        tab3 = self.tab_control.tab("Настройки")

        # Вкладка "Абоненты"
        title_frame = ctk.CTkFrame(tab1)
        title_frame.pack(fill="x", padx=10, pady=10)
        ctk.CTkLabel(title_frame, text="Управление абонентами", 
                    font=("Roboto", 16, "bold")).pack(pady=5)

        # Кнопки управления абонентами
        buttons_frame = ctk.CTkFrame(tab1)
        buttons_frame.pack(fill="x", padx=10, pady=5)

        ctk.CTkButton(buttons_frame, text="Добавить абонента",
                     command=lambda: self.create_child_window(400, 650),
                     height=35).pack(fill="x", pady=5)
        
        ctk.CTkButton(buttons_frame, text="Редактировать абонента",
                     command=self.edit_abonent,
                     height=35).pack(fill="x", pady=5)
        
        ctk.CTkButton(buttons_frame, text="Удалить абонента",
                     command=self.delete_abonent,
                     height=35).pack(fill="x", pady=5)

        # Выбор абонента
        select_frame = ctk.CTkFrame(tab1)
        select_frame.pack(fill="x", padx=10, pady=10)
        
        ctk.CTkLabel(select_frame, text="Выберите абонента:",
                    font=("Roboto", 12)).pack(pady=5)
        
        self.combobox = ctk.CTkComboBox(select_frame, width=250,
                                      font=("Roboto", 12))
        self.combobox.pack(pady=5)
        self.combobox.configure(command=self.on_combobox_select_callback)

        # Кнопки работы с данными
        data_frame = ctk.CTkFrame(tab1)
        data_frame.pack(fill="x", padx=10, pady=10)
        
        ctk.CTkButton(data_frame, text="Внести показания",
                     command=self.run_monthly_data_window,
                     height=35).pack(fill="x", pady=5)
        
        ctk.CTkButton(data_frame, text="История потребления",
                     command=self.run_consumption_history_window,
                     height=35).pack(fill="x", pady=5)

        # Вкладка "Отчеты"
        reports_title_frame = ctk.CTkFrame(tab2)
        reports_title_frame.pack(fill="x", padx=10, pady=10)
        ctk.CTkLabel(reports_title_frame, text="Генерация отчетов", 
                    font=("Roboto", 16, "bold")).pack(pady=5)

        # Кнопки для отчетов
        reports_buttons_frame = ctk.CTkFrame(tab2)
        reports_buttons_frame.pack(fill="x", padx=10, pady=5)
        
        ctk.CTkButton(reports_buttons_frame, text="📝 Создать реестры по всем абонентам",
                     command=self.generate_registries_for_all_abonents,
                     height=40,
                     font=("Roboto", 12)).pack(fill="x", pady=10)

        # Информационная панель для отчетов
        reports_info_frame = ctk.CTkFrame(tab2)
        reports_info_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        ctk.CTkLabel(reports_info_frame, text="Статус генерации отчетов",
                    font=("Roboto", 14, "bold")).pack(pady=5)
        
        self.reports_status_text = ctk.CTkTextbox(reports_info_frame,
                                                font=("Roboto", 12),
                                                wrap="word")
        self.reports_status_text.pack(fill="both", expand=True, padx=10, pady=10)

        # Вкладка "Настройки"
        settings_frame = ctk.CTkFrame(tab3)
        settings_frame.pack(fill="x", padx=10, pady=10)
        
        ctk.CTkLabel(settings_frame, text="Настройки программы",
                    font=("Roboto", 16, "bold")).pack(pady=5)
        
        ctk.CTkButton(settings_frame, text="Открыть настройки",
                     command=self.open_settings_window,
                     height=35).pack(fill="x", pady=5)

        # Информационная панель
        info_frame = ctk.CTkFrame(main_container)
        info_frame.pack(side='right', fill="both", expand=True)

        # Заголовок информационной панели
        ctk.CTkLabel(info_frame, text="Информация об абоненте",
                    font=("Roboto", 16, "bold")).pack(pady=10)

        # Поле для отображения информации об абоненте
        self.selected_abonent_info = ctk.CTkTextbox(info_frame,
                                                  font=("Roboto", 12),
                                                  wrap="word")
        self.selected_abonent_info.pack(fill="both", expand=True,
                                      padx=10, pady=10)

        # Обновляем данные в интерфейсе
        self.update_combobox()

    def update_combobox(self):
        """Обновляет список абонентов в выпадающем меню"""
        try:
            # Обновляем список абонентов
            self.list_abonent = self.load_abonents()

            if not self.list_abonent:
                print("Список абонентов пуст")
                self.combobox.configure(values=[])
                self.selected_abonent_info.delete("1.0", "end")
                return

            abonent_names = [abonent[1] for abonent in self.list_abonent]
            self.combobox.configure(values=abonent_names)

            if abonent_names:
                self.combobox.set(abonent_names[0])
                self.on_combobox_select()
        except Exception as e:
            print(f"Ошибка при обновлении combobox: {e}")

    def on_combobox_select(self, event=None):
        """Обрабатывает выбор абонента из списка"""
        try:
            selected_name = self.combobox.get()
            if not selected_name:
                return

            # Очищаем поле перед загрузкой новых данных
            self.selected_abonent_info.delete("1.0", "end")

            # Загружаем свежие данные из БД
            self.list_abonent = self.load_abonents()

            # Находим выбранного абонента
            selected_abonent = next((abonent for abonent in self.list_abonent
                                     if abonent[1] == selected_name), None)

            if not selected_abonent:
                self.selected_abonent_info.insert("1.0", "Абонент не найден")
                return

            # Основная информация об абоненте
            info = (
                f"📋 ОСНОВНАЯ ИНФОРМАЦИЯ\n"
                f"{'=' * 40}\n"
                f"🏢 Название: {selected_abonent[1]}\n\n"
                f"⚡ Электроэнергия:\n"
                f"   • Номер счетчика: {selected_abonent[2] or 'нет данных'}\n"
                f"   • Коэффициент трансформации: {selected_abonent[3] or 'нет данных'}\n\n"
                f"💧 Водоснабжение:\n"
                f"   • Номер счетчика: {selected_abonent[4] or 'нет данных'}\n"
                f"   • Водоотведение: {selected_abonent[5] or 'нет данных'}\n\n"
                f"🔥 Газоснабжение:\n"
                f"   • Номер счетчика: {selected_abonent[6] or 'нет данных'}\n\n"
            )

            # Добавляем информацию о последних месяцах с данными
            abonent_id = self.db.get_abonent_id_by_name(selected_name)
            if abonent_id:
                try:
                    last_months_data = self.db.get_last_months_data(abonent_id)
                    if last_months_data:
                        info += f"\n📊 ПОСЛЕДНИЕ ПОКАЗАНИЯ\n{'=' * 40}\n"

                        for month_data in last_months_data:
                            month, year, electricity, water, wastewater, gas = month_data
                            info += (
                                f"📅 {self.format_month(month)} {year}\n"
                                f"   • Электричество: {electricity or 'нет данных'}\n"
                                f"   • Вода: {water or 'нет данных'}\n"
                                f"   • Водоотведение: {wastewater or 'нет данных'}\n"
                                f"   • Газ: {gas or 'нет данных'}\n"
                                f"{'-' * 40}\n"
                            )
                    else:
                        info += "\n⚠️ Нет данных о потреблении\n"
                except Exception as db_error:
                    print(f"Ошибка при получении данных: {db_error}")
                    info += "\n❌ Ошибка при загрузке данных\n"

            self.selected_abonent_info.insert("1.0", info)

        except Exception as e:
            print(f"Ошибка при обработке выбора абонента: {e}")
            self.selected_abonent_info.insert("1.0", f"❌ Ошибка: {str(e)}")

    def on_combobox_select_callback(self, choice):
        """Обработчик выбора в комбобоксе"""
        self.on_combobox_select()

    def delete_abonent(self):
        """Удаляет выбранного абонента"""
        try:
            selected_name = self.combobox.get()
            if not selected_name:
                CTkMessagebox(title="Предупреждение", 
                            message="Не выбран абонент для удаления",
                            icon="warning")
                return

            confirm = CTkMessagebox(title="Подтверждение",
                                  message=f"Вы уверены, что хотите удалить абонента '{selected_name}'?",
                                  icon="question",
                                  option_1="Да",
                                  option_2="Нет")
            
            if confirm.get() == "Да":
                self.db.delete_data(selected_name)
                self.refresh_data()
                CTkMessagebox(title="Успех", 
                            message=f"Абонент '{selected_name}' удален",
                            icon="check")
        except Exception as e:
            CTkMessagebox(title="Ошибка", 
                         message=f"Ошибка при удалении абонента: {str(e)}",
                         icon="cancel")
        self.on_combobox_select()

    def refresh_data(self):
        """Полностью обновляет данные в интерфейсе"""
        try:
            # Обновляем список абонентов
            self.list_abonent = self.load_abonents()

            # Обновляем комбобокс
            if self.list_abonent:
                current_selection = self.combobox.get()
                abonent_names = [abonent[1] for abonent in self.list_abonent]
                self.combobox.configure(values=abonent_names)

                # Восстанавливаем выбор, если абонент еще существует
                if current_selection in abonent_names:
                    self.combobox.set(current_selection)
                else:
                    self.combobox.set(abonent_names[0] if abonent_names else "")

                # Обновляем информацию
                self.on_combobox_select()
            else:
                self.combobox.configure(values=[])
                self.selected_abonent_info.delete("1.0", "end")
                self.selected_abonent_info.insert("1.0", "Нет доступных абонентов")

        except Exception as e:
            CTkMessagebox(title="Ошибка", 
                         message=f"Ошибка при обновлении данных: {str(e)}",
                         icon="cancel")

    def run_monthly_data_window(self):
        """Открывает окно внесения месячных данных"""
        try:
            selected_name = self.combobox.get()
            if not selected_name:
                CTkMessagebox(title="Предупреждение", 
                            message="Сначала выберите абонента",
                            icon="warning")
                return

            abonent_id = self.db.get_abonent_id_by_name(selected_name)
            if abonent_id:
                # Создаем окно с правильной обработкой
                self.create_monthly_data_window(400, 600, abonent_id)

                # Обновляем данные
                self.refresh_data()
                self.on_combobox_select()
            else:
                CTkMessagebox(title="Ошибка", 
                            message="Не удалось определить ID абонента",
                            icon="cancel")
        except Exception as e:
            CTkMessagebox(title="Ошибка", 
                         message=f"Ошибка при открытии окна данных: {str(e)}",
                         icon="cancel")

    def edit_abonent(self):
        """Открывает окно редактирования абонента"""
        try:
            selected_name = self.combobox.get()
            if not selected_name:
                CTkMessagebox(title="Предупреждение", 
                            message="Не выбран абонент для редактирования",
                            icon="warning")
                return

            abonent_id = self.db.get_abonent_id_by_name(selected_name)
            if not abonent_id:
                CTkMessagebox(title="Ошибка", 
                            message="Не удалось определить ID абонента",
                            icon="cancel")
                return

            abonent_data = self.db.get_abonent_by_id(abonent_id)
            if not abonent_data:
                CTkMessagebox(title="Ошибка", 
                            message="Не удалось загрузить данные абонента",
                            icon="cancel")
                return

            # Создаем окно редактирования и ждем его закрытия
            edit_window = EditAbonentWindow(self.root, 400, 650, abonent_data)

            # После закрытия окна редактирования:
            self.refresh_data()  # Полностью обновляем данные
            self.combobox.set(selected_name)  # Восстанавливаем выбор абонента
            self.on_combobox_select()  # Обновляем информацию

        except Exception as e:
            CTkMessagebox(title="Ошибка", 
                         message=f"Ошибка при редактировании абонента: {str(e)}",
                         icon="cancel")

    def run_consumption_history_window(self):
        """Открывает окно истории потребления"""
        try:
            selected_name = self.combobox.get()
            if not selected_name:
                CTkMessagebox(title="Предупреждение", 
                            message="Сначала выберите абонента",
                            icon="warning")
                return

            abonent_id = self.db.get_abonent_id_by_name(selected_name)
            if abonent_id:
                self.create_consumption_history_window(900, 700, abonent_id)
            else:
                CTkMessagebox(title="Ошибка", 
                            message="Не удалось определить ID абонента",
                            icon="cancel")
        except Exception as e:
            CTkMessagebox(title="Ошибка", 
                         message=f"Ошибка при открытии окна истории: {str(e)}",
                         icon="cancel")

    def open_settings_window(self):
        """Открывает окно настроек"""
        try:
            from settings_window import SettingsWindow
            settings_window = SettingsWindow(self.root, 400, 500)
            settings_window.root.wait_window()  # Ждем закрытия окна
        except Exception as e:
            CTkMessagebox(title="Ошибка", 
                         message=f"Ошибка при открытии окна настроек: {str(e)}",
                         icon="cancel")

    def format_month(self, month_num):
        """Форматирует номер месяца в название"""
        months = [
            "Январь", "Февраль", "Март", "Апрель", "Май", "Июнь",
            "Июль", "Август", "Сентябрь", "Октябрь", "Ноябрь", "Декабрь"
        ]
        return months[month_num - 1] if 1 <= month_num <= 12 else f"Месяц {month_num}"

    def generate_registries_for_all_abonents(self):
        """Создает реестры по всем существующим абонентам"""
        try:
            # Очищаем статус
            self.reports_status_text.delete("1.0", "end")
            self.reports_status_text.insert("end", "🔄 Начинаем генерацию реестров по всем абонентам...\n\n")
            
            # Получаем всех абонентов
            all_abonents = self.db.fetch_data()
            if not all_abonents:
                self.reports_status_text.insert("end", "❌ Нет абонентов в базе данных!\n")
                return
            
            self.reports_status_text.insert("end", f"📋 Найдено абонентов: {len(all_abonents)}\n\n")
            
            # Запрашиваем месяц и год
            from tkinter import simpledialog
            import datetime
            
            current_date = datetime.datetime.now()
            month = simpledialog.askinteger("Месяц", 
                                          f"Введите месяц (1-12):", 
                                          initialvalue=current_date.month,
                                          minvalue=1, maxvalue=12)
            if month is None:
                self.reports_status_text.insert("end", "❌ Отменено пользователем\n")
                return
                
            year = simpledialog.askinteger("Год", 
                                         f"Введите год:", 
                                         initialvalue=current_date.year,
                                         minvalue=2000, maxvalue=2100)
            if year is None:
                self.reports_status_text.insert("end", "❌ Отменено пользователем\n")
                return
            
            # Получаем название месяца
            month_names = {
                1: "январь", 2: "февраль", 3: "март", 4: "апрель",
                5: "май", 6: "июнь", 7: "июль", 8: "август",
                9: "сентябрь", 10: "октябрь", 11: "ноябрь", 12: "декабрь"
            }
            month_name = month_names[month]
            
            self.reports_status_text.insert("end", f"📅 Генерируем реестры за {month_name} {year} года\n\n")
            
            # Загружаем настройки
            import os
            import json
            settings_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'settings.json')
            try:
                if os.path.exists(settings_file):
                    with open(settings_file, 'r', encoding='utf-8') as f:
                        settings = json.load(f)
                else:
                    settings = {}
            except Exception as e:
                self.reports_status_text.insert("end", f"⚠️ Ошибка при загрузке настроек: {e}\n")
                settings = {}
            
            # Получаем путь сохранения из настроек
            save_path = settings.get("save_path", r"C:\Реестры по абонентам")
            
            # Создаем подпапку с названием месяца
            month_folder_path = os.path.join(save_path, month_name)
            os.makedirs(month_folder_path, exist_ok=True)
            
            # Импортируем необходимые модули
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import re
            
            successful_count = 0
            failed_count = 0
            
            # Обрабатываем каждого абонента
            for i, abonent in enumerate(all_abonents, 1):
                abonent_id = abonent[0]
                fulname = abonent[1]
                
                self.reports_status_text.insert("end", f"📝 Обрабатываем абонента {i}/{len(all_abonents)}: {fulname}\n")
                self.reports_status_text.see("end")
                
                try:
                    # Проверяем наличие данных за указанный месяц
                    end_month_data = self.db.execute_query(
                        "SELECT * FROM monthly_data WHERE abonent_id = ? AND month = ? AND year = ?",
                        (abonent_id, month, year),
                        fetch_mode='one'
                    )
                    
                    if not end_month_data:
                        self.reports_status_text.insert("end", f"   ⚠️ Нет данных за {month_name} {year} года\n")
                        failed_count += 1
                        continue
                    
                    # Получаем данные за предыдущий месяц
                    prev_month = month - 1 if month > 1 else 12
                    prev_year = year if month > 1 else year - 1
                    
                    start_month_data = self.db.execute_query(
                        "SELECT * FROM monthly_data WHERE abonent_id = ? AND month = ? AND year = ?",
                        (abonent_id, prev_month, prev_year),
                        fetch_mode='one'
                    )
                    
                    # Создаем документ Word
                    doc = Document()
                    
                    # Настройка стилей
                    style = doc.styles['Normal']
                    style.font.name = 'Times New Roman'
                    style.font.size = Pt(12)
                    
                    # Заголовок
                    title = doc.add_paragraph()
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    title_run = title.add_run('РЕЕСТР\nвозмещения затрат за потребление электроэнергии и воды\n')
                    title_run.bold = True
                    title_run.font.size = Pt(14)
                    
                    # Период
                    period = doc.add_paragraph()
                    period.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    period.add_run(f'за {month_name} {year} г.\n').bold = True
                    
                    # Организация
                    org = doc.add_paragraph()
                    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    org.add_run(f'{fulname}\n').bold = True
                    
                    # Электроэнергия (индекс 4 в monthly_data)
                    if len(end_month_data) > 4 and end_month_data[4] is not None:
                        prev_value = float(start_month_data[4]) if start_month_data and len(start_month_data) > 4 and \
                                                                   start_month_data[4] is not None else 0.0
                        curr_value = float(end_month_data[4])
                        consumption = curr_value - prev_value
                        
                        # Получаем коэффициент трансформации из данных абонента
                        transformation_ratio = abonent[3] if len(abonent) > 3 and abonent[3] else 1
                        total_consumption = consumption * transformation_ratio
                        
                        doc.add_paragraph('1. Показания счетчика электроэнергии:', style='Normal')
                        doc.add_paragraph(f'   - на начало периода: {prev_value:.1f} кВт·ч', style='Normal')
                        doc.add_paragraph(f'   - на конец периода: {curr_value:.1f} кВт·ч', style='Normal')
                        doc.add_paragraph(f'   - итого потребление: {consumption:.1f} кВт·ч', style='Normal')
                        
                        if transformation_ratio != 1:
                            doc.add_paragraph(f'   - коэффициент трансформации: {transformation_ratio}', style='Normal')
                            doc.add_paragraph(f'   - итого потребление с учетом КТ: {total_consumption:.1f} кВт·ч', style='Normal')
                        
                        doc.add_paragraph('2. Тариф за потребленную электроэнергию: ______________ руб./кВт·ч', style='Normal')
                        doc.add_paragraph('   ИТОГО к оплате: ______________________ руб.', style='Normal')
                        doc.add_paragraph('3. Тариф за заявленную мощность: _______________ руб./кВт', style='Normal')
                        doc.add_paragraph('   Заявленная мощность: _________________ кВт', style='Normal')
                        doc.add_paragraph('   ИТОГО к оплате: ________________ руб.', style='Normal')
                        doc.add_paragraph('   ВСЕГО к оплате (п.2 + п.3): ________________ руб.', style='Normal')
                        doc.add_paragraph()
                    
                    # Вода (индекс 5)
                    if len(end_month_data) > 5 and end_month_data[5] is not None:
                        prev_value = float(start_month_data[5]) if start_month_data and len(start_month_data) > 5 and \
                                                                   start_month_data[5] is not None else 0.0
                        curr_value = float(end_month_data[5])
                        consumption = curr_value - prev_value
                        
                        doc.add_paragraph('4. Показания счетчика воды:', style='Normal')
                        doc.add_paragraph(f'   - на начало периода: {prev_value:.1f} м³', style='Normal')
                        doc.add_paragraph(f'   - на конец периода: {curr_value:.1f} м³', style='Normal')
                        doc.add_paragraph(f'   - итого потребление: {consumption:.1f} м³', style='Normal')
                        doc.add_paragraph('   Тариф: ______________ руб./м³', style='Normal')
                        doc.add_paragraph('   ИТОГО к оплате: ______________________ руб.', style='Normal')
                        doc.add_paragraph()
                    
                    # Сточные воды (индекс 6)
                    if len(end_month_data) > 6 and end_month_data[6] is not None:
                        prev_value = float(start_month_data[6]) if start_month_data and len(start_month_data) > 6 and \
                                                                   start_month_data[6] is not None else 0.0
                        curr_value = float(end_month_data[6])
                        consumption = curr_value - prev_value
                        
                        doc.add_paragraph('5. Водоотведение:', style='Normal')
                        doc.add_paragraph(f'{consumption:.1f} м³', style='Normal')
                        doc.add_paragraph('   Тариф: ______________ руб./м³', style='Normal')
                        doc.add_paragraph('   ИТОГО к оплате: ______________________ руб.', style='Normal')
                        doc.add_paragraph()
                    
                    # Газ (индекс 7)
                    if len(end_month_data) > 7 and end_month_data[7] is not None:
                        prev_value = float(start_month_data[7]) if start_month_data and len(start_month_data) > 7 and \
                                                                   start_month_data[7] is not None else 0.0
                        curr_value = float(end_month_data[7])
                        consumption = curr_value - prev_value
                        
                        doc.add_paragraph('6. Показания счетчика газа:', style='Normal')
                        doc.add_paragraph(f'   - на начало периода: {prev_value:.1f} м³', style='Normal')
                        doc.add_paragraph(f'   - на конец периода: {curr_value:.1f} м³', style='Normal')
                        doc.add_paragraph(f'   - итого потребление: {consumption:.1f} м³', style='Normal')
                        doc.add_paragraph('   Тариф: ______________ руб./м³', style='Normal')
                        doc.add_paragraph('   ИТОГО к оплате: ______________________ руб.', style='Normal')
                        doc.add_paragraph()
                    
                    # Подписи
                    signatures = settings.get("signatures", [])
                    for signature in signatures:
                        position = signature.get("position", "")
                        name = signature.get("name", "")
                        if position and name:
                            doc.add_paragraph(f'{position} {name}\t/____________/', style='Normal')
                    
                    doc.add_paragraph('Согласовано:', style='Normal')
                    doc.add_paragraph('Арендатор ___________________/________________/', style='Normal')
                    
                    # Сохраняем документ
                    safe_name = re.sub(r'[\\/*?:"<>|]', "", fulname)
                    file_name = f"{safe_name}_{month_name}_{year}_реестр.docx"
                    file_path = os.path.join(month_folder_path, file_name)
                    doc.save(file_path)
                    
                    self.reports_status_text.insert("end", f"   ✅ Реестр создан: {file_name}\n")
                    successful_count += 1
                    
                except Exception as e:
                    self.reports_status_text.insert("end", f"   ❌ Ошибка: {str(e)}\n")
                    failed_count += 1
            
            # Итоговый отчет
            self.reports_status_text.insert("end", f"\n{'='*50}\n")
            self.reports_status_text.insert("end", f"📊 ИТОГИ ГЕНЕРАЦИИ:\n")
            self.reports_status_text.insert("end", f"✅ Успешно создано: {successful_count}\n")
            self.reports_status_text.insert("end", f"❌ Ошибок: {failed_count}\n")
            self.reports_status_text.insert("end", f"📁 Папка сохранения: {month_folder_path}\n")
            
            if successful_count > 0:
                self.reports_status_text.insert("end", f"\n🎉 Генерация завершена успешно!\n")
            else:
                self.reports_status_text.insert("end", f"\n⚠️ Не удалось создать ни одного реестра\n")
            
        except Exception as e:
            self.reports_status_text.insert("end", f"❌ Критическая ошибка: {str(e)}\n")
            import traceback
            traceback.print_exc()

    def run(self):
        """Запускает главное окно"""
        try:
            self.root.mainloop()
        finally:
            # Гарантированно закрываем соединение при выходе
            if hasattr(self, 'db') and self.db:
                self.db.close_connection()


if __name__ == "__main__":
    try:
        # Проверяем и создаем структуру базы данных
        db = SqliteDB()
        db.create_table_abonent()
        db.create_table_monthly_data()
        db.close_connection()

        # Запускаем приложение
        app = Window(800, 600)
        app.run()
    except Exception as e:
        print(f"Критическая ошибка при запуске приложения: {e}")
        CTkMessagebox(title="Ошибка", 
                     message=f"Не удалось запустить приложение: {str(e)}",
                     icon="cancel")